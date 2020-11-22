from googletrans import Translator
import sys
import os
import docx
import requests
import threading
from datetime import datetime
import time

#TODO: Support 断点续传 ：resume from break-point

fname = sys.argv[1] if len(
    sys.argv) > 1 else r'G:\\'


class MyClient:
    def __init__(self):
        pass

    def get(self, url, params=None):
        return requests.request("GET", url, params=params)


def get_aligned_string(string, width):
    string = "{:{width}}".format(string, width=width)
    bts = bytes(string, 'utf-8')
    string = str(bts[0:width], encoding='utf-8', errors='backslashreplace')
    new_width = len(string) + int((width - len(string))/2)
    if new_width != 0:
        string = '{:{width}}'.format(str(string), width=new_width)
    return string


class DisplayMgr:
    def __init__(self):
        self._m = {}

    def update(self, key, value):
        self._m[key] = value

        print('--------------------')
        N = max([len(key) for key in self._m])
        for key1 in self._m:
            sign = '<' if key1 == key else ''
            aligned_key2 = get_aligned_string(key1, 80)
            print('%s %s %s' % (aligned_key2, self._m[key1], sign))


class TimeEstimator:
    def __init__(self):
        self.progress = 0
        self.t1 = datetime.now()

    def updateProgress(self, progress):
        self.progress = progress
        self.t2 = datetime.now()

    def getEstLeftTime(self):
        delta = self.t2 - self.t1
        left_delta = (1-self.progress) * delta/(self.progress+0.001)
        return left_delta


# g_TimeEstimator = TimeEstimator()
# time.sleep(1)
# g_TimeEstimator.updateProgress(0.01)
# est = g_TimeEstimator.getEstLeftTime()
# print(est)
# assert(abs(est - 1) == 0)

def MyAddRun(paragraph, text):
    run = paragraph.add_run(text)
    run.font.color.rgb = docx.shared.RGBColor(0,0,255)

def trans(fname, display_mgr):
    #translator = Translator(client=MyClient())
    translator = Translator()
    foname = fname + '-cn.docx'
    doc = docx.Document(fname)
    docdes = docx.Document(fname)

    timeEstimator = TimeEstimator()
    N = len(doc.paragraphs)
    NextTarget = 0.1
    i = 0
    while i < N:
        percentage = 1.0*i/N
        
        #if i%10==0: print(percentage)
        if percentage > NextTarget:
            outputfile = '%s-%.2f-cn.docx' % (fname, NextTarget)
            # print(outputfile)
            docdes.save(outputfile)
            NextTarget = NextTarget + 0.1

        spacer = '\n========================\n'
        spacer_short = '========================'
        subCont = doc.paragraphs[i].text_2
        j = i+1
        while len(subCont) < 4500 and j < N:
            subCont = subCont + spacer + doc.paragraphs[j].text_2
            j = j+1
        # print(fname,i,j,N)
        timeEstimator.updateProgress(percentage)
        display_mgr.update(os.path.basename(
            fname), '%5d %5d %5d %.3f %s' % (i, j, N, percentage,timeEstimator.getEstLeftTime()))
        if subCont.strip():
            # try:
            s = translator.translate(subCont, src='en', dest='zh-cn')
            ss = s.text.split(spacer_short)
            if len(ss) == j-i:
                for k in range(j-i):
                    sss = ss[k].strip()
                    if len(sss) > 0:
                        MyAddRun(docdes.paragraphs[k+i], '\n' + sss + '\n')
            else:
                print('warning translate assumption mismatch %d, %d' %
                      (len(ss), j-i))
                MyAddRun(docdes.paragraphs[j-1], '\n' + s.text + '\n')
            # except Exception as e:
            #    print('except:', e)
        i = j

    docdes.save(foname)

    timeEstimator.updateProgress(1)
    display_mgr.update(os.path.basename(
            fname), '%5d %5d %5d %.3f %s' % (N, N, N, 1.0,timeEstimator.getEstLeftTime()))

if __name__ == '__main__':
    display_mgr = DisplayMgr()
    if os.path.isfile(fname):
        trans(fname, display_mgr)
    else:
        #from multiprocessing import Process

        

        ps = []
        for filename in os.listdir(fname):
            # a file has following types
            #   is a transted file
            #   is a non translated file, and there is no correponding translated file
            is_translated_file = filename.lower().endswith('-cn.docx')
            is_docx_file = filename.lower().endswith('.docx')
            if not is_translated_file and is_docx_file:
                corres_trans_file_name = '{}.docx-cn.docx'.format(filename[:-len('.docx')])
                print(corres_trans_file_name)
                is_corres_trans_file_name_exist = os.path.isfile(corres_trans_file_name)
                if not is_corres_trans_file_name_exist:
                    foname = corres_trans_file_name
                    print(filename)
                    p = threading.Thread(target=trans, args=(
                        fname + '\\' + filename, display_mgr))
                    p.start()
                    ps.append(p)

        for p in ps:
            p.join()
