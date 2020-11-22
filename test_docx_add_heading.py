import sys
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.shared import Pt

def add_heading(file, file_new):
    document = Document(file)
    
    if 'Heading 1' in document.styles:
        document.styles['Heading 1'].delete()
    
    document_default = Document()
    #document.styles.add_style('Heading 1', document_default.styles[WD_STYLE.HEADING_1]) #Error
    #document.styles.add_style('Heading 1', document.styles.latent_styles['Heading 1']) #Error
    #document.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH) # result is not correct
    style=document.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH, True) # result is shon as heading 1but font is not correct
    #style.base_style = document_default.styles['Heading 1']
    style.font.size = Pt(30)
    #document.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH, True) # result is shon as heading 1but font is not correct
    d = document.styles.default(WD_STYLE_TYPE.PARAGRAPH)

    print('styles')
    for s in document.styles: 
        print(s.name)
    
    # 显示每段的内容
    for p in document.paragraphs:
        if p.text.lower().startswith('chapter') and p.font.size.pixel>10:
            print(p.text)
            #print(p.text_2)
            print ('change style to heading')
            #p.style = 'Heading 1'
            p.style = document.styles['Heading 1']
            #p.style = document.styles.latent_styles['Heading 1']

    # 添加段落
    paragraph = document.add_paragraph('这是新的段落内容')
    paragraph.style = document.styles['Heading 1']
    # 保存文档
    document.save(file_new)

old=r'G:\done\CUDA by Example An Introduction to General-Purpose GPU Programming by Jason Sanders, Edward Kandrot (z-lib.org).docx-cn.docx'
add_heading(old, r'G:\done\1.docx')
