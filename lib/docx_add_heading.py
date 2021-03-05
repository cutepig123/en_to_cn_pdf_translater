import sys
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.shared import Pt

def add_heading(file, file_new):
    document = Document(file)
    
    if 'Heading 1' in document.styles:
        document.styles['Heading 1'].delete()
    if 'Heading 2' in document.styles:
        document.styles['Heading 2'].delete()
    
    style=document.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH, True) # result is shon as heading 1but font is not correct
    style.font.size = Pt(30)
    style=document.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH, True) # result is shon as heading 1but font is not correct
    style.font.size = Pt(28)
    #document.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH, True) # result is shon as heading 1but font is not correct
 
    print('styles')
    for s in document.styles: 
        print(s.name)
    
    len_small_arrar={}
    starts_with_chapter_arrar={}
    is_big_font_array={}

    def maybe_chapter(i):
        if i not in len_small_arrar:
            p = document.paragraphs[i]
            if len(p.runs)>=1 and len(p.runs)<5:
                len_small_arrar[i] = len(p.runs[0].text)<100
                starts_with_chapter_arrar[i] = p.runs[0].text.lower().startswith('chapter')
                if p.runs[0].font.size is not None:
                    is_big_font_array[i] = p.runs[0].font.size.pt>=20
                else:
                    is_big_font_array[i] = False
            else:
                len_small_arrar[i] = False
                starts_with_chapter_arrar[i] = False
                is_big_font_array[i] = False

        return [len_small_arrar[i], starts_with_chapter_arrar[i], is_big_font_array[i]]
    

    # 显示每段的内容
    nP = len(document.paragraphs)
    nChap = 0
    for i in range(nP-1):
        if i%10==0: print(i, nChap, 1.0*i/nP)
        [len_small, starts_with_chapter, is_big_font] = maybe_chapter(i)
        if len_small and starts_with_chapter:
            nChap = nChap +1
            #print(p.text)
            #print(p.text_2)
            print ('change style to heading')
            document.paragraphs[i].style = document.styles['Heading 1']
        elif len_small and is_big_font:
            nChap = nChap +1
            #print(p.text)
            #print(p.text_2)
            print ('change style to heading 2')
            document.paragraphs[i].style = document.styles['Heading 2']
            
            
    # 添加段落
    paragraph = document.add_paragraph('这是新的段落内容')
    paragraph.style = document.styles['Heading 1']
    # 保存文档
    document.save(file_new)
    return True
