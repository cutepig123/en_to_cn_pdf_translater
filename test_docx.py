from docx import Document
# 从文件创建文档对象
document = Document(r'G:\doc\done\Fetch-and-add.docx')
# 显示每段的内容
for p in document.paragraphs:
    print(p.text)
    print(p.text_2)
# 添加段落
document.add_paragraph('这是新的段落内容')
# 保存文档
document.save('demo.docx')