from googletrans import Translator
import sys, os
import docx
import requests

fname = sys.argv[1] if len(
    sys.argv) > 1 else r'G:\\'

class MyClient:
    def __init__(self):
        pass

    def get(self, url, params = None):
        return requests.request("GET", url, params = params)

class DisplayMgr:
    def __init__(self):
        self._m = {}

    def update(self, key, value):
        self._m[key] = value

        for key in self._m:
            print(key, value)
            
def trans(fname):
    #translator = Translator(client=MyClient())
    translator = Translator()
    foname = fname + '-cn.docx'
    doc = docx.Document(fname)
    docdes = docx.Document(fname)

    N = len(doc.paragraphs)
    NextTarget = 0.1
    i = 0
    while i<N:
        percentage = 1.0*i/N
        if i%10==0: print(percentage)
        if percentage>NextTarget:
            outputfile = '%s-%.2f-cn.docx'%(fname, NextTarget)
            print(outputfile)
            docdes.save(outputfile)
            NextTarget = NextTarget + 0.1

        spacer = '\n========================\n'
        spacer_short = '========================'
        subCont = doc.paragraphs[i].text
        j = i+1
        while len(subCont)<4500 and j<N:
            subCont = subCont + spacer + doc.paragraphs[j].text
            j = j+1
        print(fname,i,j,N)
        if subCont.strip():
            #try:
            s = translator.translate(subCont, src='en', dest='zh-cn')
            ss = s.text.split(spacer_short)
            if len(ss)==j-i:
                for k in range(j-i):
                    sss = ss[k].strip()
                    if len(sss)>0:
                        docdes.paragraphs[k+i].add_run('\n' + sss + '\n')
            else:
                print ('warning translate assumption mismatch %d, %d'%(len(ss), j-i))
                docdes.paragraphs[j-1].add_run('\n' + s.text + '\n')
            #except Exception as e:
            #    print('except:', e)
        i = j

    docdes.save(foname)

if __name__ == '__main__':
    if os.path.isfile(fname):
        trans(fname)
    else:
        from multiprocessing import Process

        ps=[]
        for filename in os.listdir(fname):
            foname = fname + '\\' + filename + '-cn.docx'
            if filename.lower().endswith('.docx') and not filename.lower().endswith('-cn.docx') and not os.path.isfile(foname):
                print(filename)
                p = Process(target=trans, args=(fname + '\\' + filename,))
                p.start()
                ps.append(p)

        for p in ps:
            p.join()